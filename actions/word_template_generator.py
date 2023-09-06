import datetime
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches

output_dir = "output_data"
file_counter = 1
#template_dir = "note_templates/template.docx"


def generate_cumulo_template(doc, data_row, currency, producto):
    from actions.file_actions import translate_month_to_spanish, get_current_month, format_date
    # Add texto exclusion
    monto = "3.000.000.000"
    exclusion_paragraph = doc.add_paragraph(("      Por la presente se informa la exclusión del Prestatario indicado a continuación, de la póliza de Seguro de Vida Colectivo para Cancelación de Deudas, por superar el capital de Gs. {monto}, establecido como cúmulo máximo por Asegurado.").format(monto=monto))
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph(("      La operación corresponde a la planilla de {producto} en moneda {moneda} del mes de {mes}.").format(producto=producto, moneda=currency, mes=translate_month_to_spanish(get_current_month())))
    exclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    exclusion_paragraph.runs[0].font.name = 'Arial'

    # Add a table asegurado
    asegurado_table = doc.add_table(rows=3, cols=2)
    asegurado_table.style = 'Table Grid'
    nombre_cells = asegurado_table.rows[0].cells
    nombre_cells[0].text = 'Asegurado:'
    nombre_cells[1].text = data_row[0]
    documento_cells = asegurado_table.rows[1].cells
    documento_cells[0].text = 'Documento:'
    documento_cells[1].text = str(data_row[1])
    nacimiento_cells = asegurado_table.rows[2].cells
    nacimiento_cells[0].text = 'Fecha de Nacimiento:'
    nacimiento_cells[1].text = format_date(data_row[2])
    for row in asegurado_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'

    # Add a section break
    doc.add_paragraph(" ")
    
    # Add a table operacion
    operacion_table = doc.add_table(rows=2, cols=4)
    operacion_table.style = 'Table Grid'
    hdr_cells = operacion_table.rows[0].cells
    hdr_cells[0].text = 'Nro. Operación'
    hdr_cells[1].text = 'Monto'
    hdr_cells[2].text = 'Costo'
    hdr_cells[3].text = 'Fecha Vencimiento'
    for cell in operacion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    row_cells = operacion_table.rows[1].cells
    row_cells[0].text = str(data_row[4])
    if currency == "guaraníes":
        row_cells[1].text = str("{:,.0f}".format(data_row[6])).replace(",", ".")
        row_cells[2].text = str("{:,.0f}".format(data_row[8])).replace(",", ".")
    elif currency == "dólares americanos":
        row_cells[1].text = str("{:,.2f}".format(data_row[6])).replace(".", "x").replace(",", ".").replace("x", ",")
        row_cells[2].text = str("{:,.2f}".format(data_row[8])).replace(".", "x").replace(",", ".").replace("x", ",")
    row_cells[3].text = format_date(data_row[10]) 
    for row in operacion_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'
    
    # Definir el color gris (RGB: 192, 192, 192)
    gray_color = RGBColor(240, 240, 240)

    # Agregar fondo de color gris a la celda
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}" />')
        tcPr.append(shading_elm)
   
    # Add a section break
    doc.add_paragraph(" ")


def generate_fallecimiento_template(doc, data_row, currency, producto):
    from actions.file_actions import translate_month_to_spanish, get_current_month, format_date
    # Add texto exclusion
    exclusion_paragraph = doc.add_paragraph("       Por la presente se informa la exclusión del Prestatario indicado a continuación, de la póliza de Seguro de Vida Colectivo para Cancelación de Deudas, por fallecimiento.")
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph(("       La operación corresponde a la planilla de {producto} en moneda {moneda} del mes de {mes}.").format(producto=producto, moneda=currency, mes=translate_month_to_spanish(get_current_month())))
    exclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    exclusion_paragraph.runs[0].font.name = 'Arial'

    # Add a table operacion
    operacion_table = doc.add_table(rows=2, cols=8)
    operacion_table.style = 'Table Grid'
    hdr_cells = operacion_table.rows[0].cells
    hdr_cells[0].text = 'Nombre del cliente'
    hdr_cells[1].text = 'Nro. Documento'
    hdr_cells[2].text = 'Fecha Nacimiento'
    hdr_cells[3].text = 'Capital Asegurado'
    hdr_cells[4].text = 'Costo del Seguro'
    hdr_cells[5].text = 'Fecha Inicio'
    hdr_cells[6].text = 'Fecha Vencimiento'
    hdr_cells[7].text = 'Nro. Operación'
    for cell in operacion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    row_cells = operacion_table.rows[1].cells
    row_cells[0].text = data_row[0]
    row_cells[1].text = str(data_row[1])
    row_cells[2].text = format_date(data_row[2])
    if currency == "guaraníes":
        row_cells[3].text = str("{:,.0f}".format(data_row[6])).replace(",", ".")
        row_cells[4].text = str("{:,.0f}".format(data_row[8])).replace(",", ".")
    elif currency == "dólares americanos":
        row_cells[3].text = str("{:,.2f}".format(data_row[6])).replace(".", "x").replace(",", ".").replace("x", ",")
        row_cells[4].text = str("{:,.2f}".format(data_row[8])).replace(".", "x").replace(",", ".").replace("x", ",")
    row_cells[5].text = format_date(data_row[9])
    row_cells[6].text = format_date(data_row[10]) 
    row_cells[7].text = str(data_row[4])
    for row in operacion_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'

    # Definir el color gris (RGB: 192, 192, 192)
    gray_color = RGBColor(240, 240, 240)

    # Agregar fondo de color gris a la celda
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}" />')
        tcPr.append(shading_elm)
   
    # Add a section break
    doc.add_paragraph(" ")
    

def generate_falta_ds_template(doc, data_row, currency, producto):
    from actions.file_actions import translate_month_to_spanish, get_current_month, format_date
    # Add texto exclusion
    exclusion_paragraph = doc.add_paragraph("       Por la presente se informa la exclusión del Prestatario indicado a continuación, de la póliza de Seguro de Vida Colectivo para Cancelación de Deudas, debido a que no se ha presentado el formulario de Declaración de Salud.")
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph(("       La operación corresponde a la planilla de {producto} en moneda {moneda} del mes de {mes}.").format(producto=producto, moneda=currency, mes=translate_month_to_spanish(get_current_month())))
    exclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    exclusion_paragraph.runs[0].font.name = 'Arial'

    # Add a table operacion
    operacion_table = doc.add_table(rows=2, cols=8)
    operacion_table.style = 'Table Grid'
    hdr_cells = operacion_table.rows[0].cells
    hdr_cells[0].text = 'Nombre del cliente'
    hdr_cells[1].text = 'Nro. Documento'
    hdr_cells[2].text = 'Fecha Nacimiento'
    hdr_cells[3].text = 'Capital Asegurado'
    hdr_cells[4].text = 'Costo del Seguro'
    hdr_cells[5].text = 'Fecha Inicio'
    hdr_cells[6].text = 'Fecha Vencimiento'
    hdr_cells[7].text = 'Nro. Operación'
    for cell in operacion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    row_cells = operacion_table.rows[1].cells
    row_cells[0].text = data_row[0]
    row_cells[1].text = str(data_row[1])
    row_cells[2].text = format_date(data_row[2])
    if currency == "guaraníes":
        row_cells[3].text = str("{:,.0f}".format(data_row[6])).replace(",", ".")
        row_cells[4].text = str("{:,.0f}".format(data_row[8])).replace(",", ".")
    elif currency == "dólares americanos":
        row_cells[3].text = str("{:,.2f}".format(data_row[6])).replace(".", "x").replace(",", ".").replace("x", ",")
        row_cells[4].text = str("{:,.2f}".format(data_row[8])).replace(".", "x").replace(",", ".").replace("x", ",")
    row_cells[5].text = format_date(data_row[9])
    row_cells[6].text = format_date(data_row[10]) 
    row_cells[7].text = str(data_row[4])
    for row in operacion_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'

    # Definir el color gris (RGB: 192, 192, 192)
    gray_color = RGBColor(240, 240, 240)

    # Agregar fondo de color gris a la celda
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}" />')
        tcPr.append(shading_elm)
   
    # Add a section break
    doc.add_paragraph(" ")
    

def generate_ds_incompleta_template(doc, data_row, currency, producto):
    from actions.file_actions import translate_month_to_spanish, get_current_month, format_date
    # Add texto exclusion
    exclusion_paragraph = doc.add_paragraph("       Por la presente se informa la exclusión del Prestatario indicado a continuación, de la póliza de Seguro de Vida Colectivo para Cancelación de Deudas, por presentar incompleta su Declaración de Salud.")
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph(("       La operación corresponde a la planilla de {producto} en moneda {moneda} del mes de {mes}.").format(producto=producto, moneda=currency, mes=translate_month_to_spanish(get_current_month())))
    exclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    exclusion_paragraph.runs[0].font.name = 'Arial'

    # Add a table operacion
    operacion_table = doc.add_table(rows=2, cols=8)
    operacion_table.style = 'Table Grid'
    hdr_cells = operacion_table.rows[0].cells
    hdr_cells[0].text = 'Nombre del cliente'
    hdr_cells[1].text = 'Nro. Documento'
    hdr_cells[2].text = 'Fecha Nacimiento'
    hdr_cells[3].text = 'Capital Asegurado'
    hdr_cells[4].text = 'Costo del Seguro'
    hdr_cells[5].text = 'Fecha Inicio'
    hdr_cells[6].text = 'Fecha Vencimiento'
    hdr_cells[7].text = 'Nro. Operación'
    for cell in operacion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    row_cells = operacion_table.rows[1].cells
    row_cells[0].text = data_row[0]
    row_cells[1].text = str(data_row[1])
    row_cells[2].text = format_date(data_row[2])
    if currency == "guaraníes":
        row_cells[3].text = str("{:,.0f}".format(data_row[6])).replace(",", ".")
        row_cells[4].text = str("{:,.0f}".format(data_row[8])).replace(",", ".")
    elif currency == "dólares americanos":
        row_cells[3].text = str("{:,.2f}".format(data_row[6])).replace(".", "x").replace(",", ".").replace("x", ",")
        row_cells[4].text = str("{:,.2f}".format(data_row[8])).replace(".", "x").replace(",", ".").replace("x", ",")
    row_cells[5].text = format_date(data_row[9])
    row_cells[6].text = format_date(data_row[10]) 
    row_cells[7].text = str(data_row[4])
    for row in operacion_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'

    # Definir el color gris (RGB: 192, 192, 192)
    gray_color = RGBColor(240, 240, 240)

    # Agregar fondo de color gris a la celda
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}" />')
        tcPr.append(shading_elm)
   
    # Add a section break
    doc.add_paragraph(" ")
    despedida_paragraph = doc.add_paragraph("       En tal sentido y hasta la regularización de la Declaración de Salud, el riesgo queda excluido de la póliza. Se adjunta copia del formulario.")
    despedida_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY 
    despedida_paragraph.runs[0].font.name = 'Arial'


def generate_sin_capital_template(doc, data_row, currency, producto):
    from actions.file_actions import translate_month_to_spanish, get_current_month, format_date
    # Add texto exclusion
    exclusion_paragraph = doc.add_paragraph("       Por la presente se informa la exclusión del Prestatario indicado a continuación, de la póliza de Seguro de Vida Colectivo para Cancelación de Deudas, debido a que no cuenta con capital en la planilla.")
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph(("       La operación corresponde a la planilla de {producto} en moneda {moneda} del mes de {mes}.").format(producto=producto, moneda=currency, mes=translate_month_to_spanish(get_current_month())))
    exclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    exclusion_paragraph.runs[0].font.name = 'Arial'

    # Add a table operacion
    operacion_table = doc.add_table(rows=2, cols=7)
    operacion_table.style = 'Table Grid'
    hdr_cells = operacion_table.rows[0].cells
    hdr_cells[0].text = 'Nombre del cliente'
    hdr_cells[1].text = 'Nro. Documento'
    hdr_cells[2].text = 'Fecha Nacimiento'
    hdr_cells[3].text = 'Capital Asegurado'
    hdr_cells[4].text = 'Costo del Seguro'
    hdr_cells[5].text = 'Fecha Vencimiento'
    hdr_cells[6].text = 'Nro. Operación'
    for cell in operacion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    row_cells = operacion_table.rows[1].cells
    row_cells[0].text = data_row[0]
    row_cells[1].text = str(data_row[1])
    row_cells[2].text = format_date(data_row[2])
    if currency == "guaraníes":
        row_cells[3].text = str("{:,.0f}".format(data_row[6])).replace(",", ".")
        row_cells[4].text = str("{:,.0f}".format(data_row[8])).replace(",", ".")
    elif currency == "dólares americanos":
        row_cells[3].text = str("{:,.2f}".format(data_row[6])).replace(".", "x").replace(",", ".").replace("x", ",")
        row_cells[4].text = str("{:,.2f}".format(data_row[8])).replace(".", "x").replace(",", ".").replace("x", ",")
    row_cells[5].text = format_date(data_row[10]) 
    row_cells[6].text = str(data_row[4])
    for row in operacion_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'

    # Definir el color gris (RGB: 192, 192, 192)
    gray_color = RGBColor(240, 240, 240)

    # Agregar fondo de color gris a la celda
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}" />')
        tcPr.append(shading_elm)
   
    # Add a section break
    doc.add_paragraph(" ")


def generate_mora_template(doc, data_row, currency, producto):
    from actions.file_actions import translate_month_to_spanish, get_current_month, format_date
    # Add texto exclusion
    exclusion_paragraph = doc.add_paragraph("       Por la presente se informa la exclusión del Prestatario indicado a continuación, de la póliza de Seguro de Vida Colectivo para Cancelación de Deudas, debido a que posee operaciones en mora.")
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph(("       La operación corresponde a la planilla de {producto} en moneda {moneda} del mes de {mes}.").format(producto=producto, moneda=currency, mes=translate_month_to_spanish(get_current_month())))
    exclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    exclusion_paragraph.runs[0].font.name = 'Arial'

    # Add a table operacion
    operacion_table = doc.add_table(rows=2, cols=8)
    operacion_table.style = 'Table Grid'
    hdr_cells = operacion_table.rows[0].cells
    hdr_cells[0].text = 'Nombre del cliente'
    hdr_cells[1].text = 'Nro. Documento'
    hdr_cells[2].text = 'Fecha Nacimiento'
    hdr_cells[3].text = 'Capital Asegurado'
    hdr_cells[4].text = 'Costo del Seguro'
    hdr_cells[5].text = 'Fecha Inicio'
    hdr_cells[6].text = 'Fecha Vencimiento'
    hdr_cells[7].text = 'Nro. Operación'
    for cell in operacion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    row_cells = operacion_table.rows[1].cells
    row_cells[0].text = data_row[0]
    row_cells[1].text = str(data_row[1])
    row_cells[2].text = format_date(data_row[2])
    if currency == "guaraníes":
        row_cells[3].text = str("{:,.0f}".format(data_row[6])).replace(",", ".")
        row_cells[4].text = str("{:,.0f}".format(data_row[8])).replace(",", ".")
    elif currency == "dólares americanos":
        row_cells[3].text = str("{:,.2f}".format(data_row[6])).replace(".", "x").replace(",", ".").replace("x", ",")
        row_cells[4].text = str("{:,.2f}".format(data_row[8])).replace(".", "x").replace(",", ".").replace("x", ",")
    row_cells[5].text = format_date(data_row[9]) 
    row_cells[6].text = format_date(data_row[10]) 
    row_cells[7].text = str(data_row[4])
    for row in operacion_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'

    # Definir el color gris (RGB: 192, 192, 192)
    gray_color = RGBColor(240, 240, 240)

    # Agregar fondo de color gris a la celda
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}" />')
        tcPr.append(shading_elm)
   
    # Add a section break
    doc.add_paragraph(" ")    


def generate_edad_template(doc, data_row, currency, producto):
    from actions.file_actions import translate_month_to_spanish, get_current_month, format_date
    # Add texto exclusion
    exclusion_paragraph = doc.add_paragraph("       Por la presente se informa la exclusión del Prestatario indicado a continuación, de la póliza de Seguro de Vida Colectivo para Cancelación de Deudas, por superar la edad límite de 75 años.")
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph(("       La operación corresponde a la planilla de {producto} en moneda {moneda} del mes de {mes}.").format(producto=producto, moneda=currency, mes=translate_month_to_spanish(get_current_month())))
    exclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    exclusion_paragraph.runs[0].font.name = 'Arial'

    # Add a table operacion
    operacion_table = doc.add_table(rows=2, cols=7)
    operacion_table.style = 'Table Grid'
    hdr_cells = operacion_table.rows[0].cells
    hdr_cells[0].text = 'Nombre del cliente'
    hdr_cells[1].text = 'Nro. Documento'
    hdr_cells[2].text = 'Fecha Nacimiento'
    hdr_cells[3].text = 'Capital Asegurado'
    hdr_cells[4].text = 'Costo del Seguro'
    hdr_cells[5].text = 'Fecha Vencimiento'
    hdr_cells[6].text = 'Nro. Operación'
    for cell in operacion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    row_cells = operacion_table.rows[1].cells
    row_cells[0].text = data_row[0]
    row_cells[1].text = str(data_row[1])
    row_cells[2].text = format_date(data_row[2])
    if currency == "guaraníes":
        row_cells[3].text = str("{:,.0f}".format(data_row[6])).replace(",", ".")
        row_cells[4].text = str("{:,.0f}".format(data_row[8])).replace(",", ".")
    elif currency == "dólares americanos":
        row_cells[3].text = str("{:,.2f}".format(data_row[6])).replace(".", "x").replace(",", ".").replace("x", ",")
        row_cells[4].text = str("{:,.2f}".format(data_row[8])).replace(".", "x").replace(",", ".").replace("x", ",")
    row_cells[5].text = format_date(data_row[10]) 
    row_cells[6].text = str(data_row[4])
    for row in operacion_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'

    # Definir el color gris (RGB: 192, 192, 192)
    gray_color = RGBColor(240, 240, 240)

    # Agregar fondo de color gris a la celda
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}" />')
        tcPr.append(shading_elm)
   
    # Add a section break
    doc.add_paragraph(" ")


def generate_operacion_vencida_template(doc, data_row, currency, producto):
    from actions.file_actions import translate_month_to_spanish, get_current_month, format_date
    # Add texto exclusion
    exclusion_paragraph = doc.add_paragraph("       Por la presente se informa la exclusión del Prestatario indicado a continuación, de la póliza de Seguro de Vida Colectivo para Cancelación de Deudas, debido a que sus operaciones se encuentran vencidas.")
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph(("       La operación corresponde a la planilla de {producto} en moneda {moneda} del mes de {mes}.").format(producto=producto, moneda=currency, mes=translate_month_to_spanish(get_current_month())))
    exclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    exclusion_paragraph.runs[0].font.name = 'Arial'

    # Add a table operacion
    operacion_table = doc.add_table(rows=2, cols=7)
    operacion_table.style = 'Table Grid'
    hdr_cells = operacion_table.rows[0].cells
    hdr_cells[0].text = 'Nombre del cliente'
    hdr_cells[1].text = 'Nro. Documento'
    hdr_cells[2].text = 'Fecha Nacimiento'
    hdr_cells[3].text = 'Capital Asegurado'
    hdr_cells[4].text = 'Costo del Seguro'
    hdr_cells[5].text = 'Fecha Vencimiento'
    hdr_cells[6].text = 'Nro. Operación'
    for cell in operacion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    row_cells = operacion_table.rows[1].cells
    row_cells[0].text = data_row[0]
    row_cells[1].text = str(data_row[1])
    row_cells[2].text = format_date(data_row[2])
    if currency == "guaraníes":
        row_cells[3].text = str("{:,.0f}".format(data_row[6])).replace(",", ".")
        row_cells[4].text = str("{:,.0f}".format(data_row[8])).replace(",", ".")
    elif currency == "dólares americanos":
        row_cells[3].text = str("{:,.2f}".format(data_row[6])).replace(".", "x").replace(",", ".").replace("x", ",")
        row_cells[4].text = str("{:,.2f}".format(data_row[8])).replace(".", "x").replace(",", ".").replace("x", ",")
    row_cells[5].text = format_date(data_row[10]) 
    row_cells[6].text = str(data_row[4])
    for row in operacion_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'

    # Definir el color gris (RGB: 192, 192, 192)
    gray_color = RGBColor(240, 240, 240)

    # Agregar fondo de color gris a la celda
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}" />')
        tcPr.append(shading_elm)
   
    # Add a section break
    doc.add_paragraph(" ")


def generate_policita_suscripcion_template(doc, data_row, currency, producto):
    from actions.file_actions import translate_month_to_spanish, get_current_month, format_date
    # Add texto exclusion
    exclusion_paragraph = doc.add_paragraph("       Por la presente se informa la exclusión del Prestatario indicado a continuación, de la póliza de Seguro de Vida Colectivo para Cancelación de Deudas, por no adecuarse a la Política de Suscripción de la Compañía.")
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph(("       La operación corresponde a la planilla de {producto} en moneda {moneda} del mes de {mes}.").format(producto=producto, moneda=currency, mes=translate_month_to_spanish(get_current_month())))
    exclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    exclusion_paragraph.runs[0].font.name = 'Arial'

    # Add a table operacion
    operacion_table = doc.add_table(rows=2, cols=8)
    operacion_table.style = 'Table Grid'
    hdr_cells = operacion_table.rows[0].cells
    hdr_cells[0].text = 'Nombre del cliente'
    hdr_cells[1].text = 'Nro. Documento'
    hdr_cells[2].text = 'Fecha Nacimiento'
    hdr_cells[3].text = 'Capital Asegurado'
    hdr_cells[4].text = 'Costo del Seguro'
    hdr_cells[5].text = 'Fecha Inicio'
    hdr_cells[6].text = 'Fecha Vencimiento'
    hdr_cells[7].text = 'Nro. Operación'
    for cell in operacion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    row_cells = operacion_table.rows[1].cells
    row_cells[0].text = data_row[0]
    row_cells[1].text = str(data_row[1])
    row_cells[2].text = format_date(data_row[2])
    if currency == "guaraníes":
        row_cells[3].text = str("{:,.0f}".format(data_row[6])).replace(",", ".")
        row_cells[4].text = str("{:,.0f}".format(data_row[8])).replace(",", ".")
    elif currency == "dólares americanos":
        row_cells[3].text = str("{:,.2f}".format(data_row[6])).replace(".", "x").replace(",", ".").replace("x", ",")
        row_cells[4].text = str("{:,.2f}".format(data_row[8])).replace(".", "x").replace(",", ".").replace("x", ",")
    row_cells[5].text = format_date(data_row[9]) 
    row_cells[6].text = format_date(data_row[10]) 
    row_cells[7].text = str(data_row[4])
    for row in operacion_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'

    # Definir el color gris (RGB: 192, 192, 192)
    gray_color = RGBColor(240, 240, 240)

    # Agregar fondo de color gris a la celda
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}" />')
        tcPr.append(shading_elm)
   
    # Add a section break
    doc.add_paragraph(" ")


def generate_anulado_template(doc, data_row, currency, producto):
    from actions.file_actions import translate_month_to_spanish, get_current_month, format_date
    # Add texto exclusion
    exclusion_paragraph = doc.add_paragraph("       Por la presente se informa la exclusión del Prestatario indicado a continuación, de la póliza de Seguro de Vida Colectivo para Cancelación de Deudas, debido a que poseen operaciones que fueron remitidas para anular previamente.")
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph(("       La operación corresponde a la planilla de {producto} en moneda {moneda} del mes de {mes}.").format(producto=producto, moneda=currency, mes=translate_month_to_spanish(get_current_month())))
    exclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    exclusion_paragraph.runs[0].font.name = 'Arial'

    # Add a table operacion
    operacion_table = doc.add_table(rows=2, cols=7)
    operacion_table.style = 'Table Grid'
    hdr_cells = operacion_table.rows[0].cells
    hdr_cells[0].text = 'Nombre del cliente'
    hdr_cells[1].text = 'Nro. Documento'
    hdr_cells[2].text = 'Fecha Nacimiento'
    hdr_cells[3].text = 'Capital Asegurado'
    hdr_cells[4].text = 'Costo del Seguro'
    hdr_cells[5].text = 'Fecha Vencimiento'
    hdr_cells[6].text = 'Nro. Operación'
    for cell in operacion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    row_cells = operacion_table.rows[1].cells
    row_cells[0].text = data_row[0]
    row_cells[1].text = str(data_row[1])
    row_cells[2].text = format_date(data_row[2])
    if currency == "guaraníes":
        row_cells[3].text = str("{:,.0f}".format(data_row[6])).replace(",", ".")
        row_cells[4].text = str("{:,.0f}".format(data_row[8])).replace(",", ".")
    elif currency == "dólares americanos":
        row_cells[3].text = str("{:,.2f}".format(data_row[6])).replace(".", "x").replace(",", ".").replace("x", ",")
        row_cells[4].text = str("{:,.2f}".format(data_row[8])).replace(".", "x").replace(",", ".").replace("x", ",")
    row_cells[5].text = format_date(data_row[10]) 
    row_cells[6].text = str(data_row[4])
    for row in operacion_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'

    # Definir el color gris (RGB: 192, 192, 192)
    gray_color = RGBColor(240, 240, 240)

    # Agregar fondo de color gris a la celda
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}" />')
        tcPr.append(shading_elm)
   
    # Add a section break
    doc.add_paragraph(" ")


def generate_diferencia_ds_template(doc, data_row, currency, producto):
    from actions.file_actions import translate_month_to_spanish, get_current_month, format_date
    # Add texto exclusion
    exclusion_paragraph = doc.add_paragraph("       Por la presente se informa la exclusión del Prestatario indicado a continuación, de la póliza de Seguro de Vida Colectivo para Cancelación de Deudas, por presentar diferencias entre los datos consignados en la Declaración de Salud y la planilla de Seguros.")
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph(("       La operación corresponde a la planilla de {producto} en moneda {moneda} del mes de {mes}.").format(producto=producto, moneda=currency, mes=translate_month_to_spanish(get_current_month())))
    exclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    exclusion_paragraph.runs[0].font.name = 'Arial'

    # Add a table operacion
    operacion_table = doc.add_table(rows=2, cols=8)
    operacion_table.style = 'Table Grid'
    hdr_cells = operacion_table.rows[0].cells
    hdr_cells[0].text = 'Nombre del cliente'
    hdr_cells[1].text = 'Nro. Documento'
    hdr_cells[2].text = 'Fecha Nacimiento'
    hdr_cells[3].text = 'Capital Asegurado'
    hdr_cells[4].text = 'Costo del Seguro'
    hdr_cells[5].text = 'Fecha Inicio'
    hdr_cells[6].text = 'Fecha Vencimiento'
    hdr_cells[7].text = 'Nro. Operación'
    for cell in operacion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    row_cells = operacion_table.rows[1].cells
    row_cells[0].text = data_row[0]
    row_cells[1].text = str(data_row[1])
    row_cells[2].text = format_date(data_row[2])
    if currency == "guaraníes":
        row_cells[3].text = str("{:,.0f}".format(data_row[6])).replace(",", ".")
        row_cells[4].text = str("{:,.0f}".format(data_row[8])).replace(",", ".")
    elif currency == "dólares americanos":
        row_cells[3].text = str("{:,.2f}".format(data_row[6])).replace(".", "x").replace(",", ".").replace("x", ",")
        row_cells[4].text = str("{:,.2f}".format(data_row[8])).replace(".", "x").replace(",", ".").replace("x", ",")
    row_cells[5].text = format_date(data_row[9]) 
    row_cells[6].text = format_date(data_row[10]) 
    row_cells[7].text = str(data_row[4])
    for row in operacion_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'

    # Definir el color gris (RGB: 192, 192, 192)
    gray_color = RGBColor(240, 240, 240)

    # Agregar fondo de color gris a la celda
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}" />')
        tcPr.append(shading_elm)
   
    # Add a section break
    doc.add_paragraph(" ")
    despedida_paragraph = doc.add_paragraph("       En tal sentido y hasta la regularización de la Declaración de Salud, el riesgo queda excluido de la póliza. Se adjunta copia del formulario.")
    despedida_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY 
    despedida_paragraph.runs[0].font.name = 'Arial'


def generate_informacion_adicional_template(doc, data_row, currency, producto):
    from actions.file_actions import translate_month_to_spanish, get_current_month, format_date
    # Add texto exclusion
    exclusion_paragraph = doc.add_paragraph("       Por la presente se informa la exclusión del Prestatario indicado a continuación, de la póliza de Seguro de Vida Colectivo para Cancelación de Deudas, por motivo de solicitud de ingormación adicional con respecto a lo declarado en los ítems de la Declaración de Salud.")
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph(("       La operación corresponde a la planilla de {producto} en moneda {moneda} del mes de {mes}.").format(producto=producto, moneda=currency, mes=translate_month_to_spanish(get_current_month())))
    exclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    exclusion_paragraph.runs[0].font.name = 'Arial'

    # Add a table operacion
    operacion_table = doc.add_table(rows=2, cols=8)
    operacion_table.style = 'Table Grid'
    hdr_cells = operacion_table.rows[0].cells
    hdr_cells[0].text = 'Nombre del cliente'
    hdr_cells[1].text = 'Nro. Documento'
    hdr_cells[2].text = 'Fecha Nacimiento'
    hdr_cells[3].text = 'Capital Asegurado'
    hdr_cells[4].text = 'Costo del Seguro'
    hdr_cells[5].text = 'Fecha Inicio'
    hdr_cells[6].text = 'Fecha Vencimiento'
    hdr_cells[7].text = 'Nro. Operación'
    for cell in operacion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    row_cells = operacion_table.rows[1].cells
    row_cells[0].text = data_row[0]
    row_cells[1].text = str(data_row[1])
    row_cells[2].text = format_date(data_row[2])
    if currency == "guaraníes":
        row_cells[3].text = str("{:,.0f}".format(data_row[6])).replace(",", ".")
        row_cells[4].text = str("{:,.0f}".format(data_row[8])).replace(",", ".")
    elif currency == "dólares americanos":
        row_cells[3].text = str("{:,.2f}".format(data_row[6])).replace(".", "x").replace(",", ".").replace("x", ",")
        row_cells[4].text = str("{:,.2f}".format(data_row[8])).replace(".", "x").replace(",", ".").replace("x", ",")
    row_cells[5].text = format_date(data_row[9]) 
    row_cells[6].text = format_date(data_row[10]) 
    row_cells[7].text = str(data_row[4])
    for row in operacion_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'

    # Definir el color gris (RGB: 192, 192, 192)
    gray_color = RGBColor(240, 240, 240)

    # Agregar fondo de color gris a la celda
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}" />')
        tcPr.append(shading_elm)
   
    # Add a section break
    doc.add_paragraph(" ")
    despedida_paragraph = doc.add_paragraph("""       Se solicita el Informe Médico que detalle los siguientes ítems:
•	Antigüedad de la dolencia (o tiempo en el que fue descubierta).
•	Tratamiento médico prescripto y su duración y respuesta al tratamiento.
•	Condición actual del solicitante con respecto la dolencia.
""")
    despedida_paragraph.runs[0].font.name = 'Arial'
    despedida_paragraph = doc.add_paragraph("       En tal sentido y hasta la regularización de la Declaración de Salud, el riesgo queda excluido de la póliza. Se adjunta copia del formulario.")
    despedida_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY 
    despedida_paragraph.runs[0].font.name = 'Arial'


def generate_exclusiones_previas_template(doc, data_row, currency, producto):
    from actions.file_actions import translate_month_to_spanish, get_current_month, format_date
    # Add texto exclusion
    exclusion_paragraph = doc.add_paragraph("       Por la presente se informa la exclusión del Prestatario indicado a continuación, de la póliza de Seguro de Vida Colectivo para Cancelación de Deudas, debido a que sus operaciones han sido excluidas previamente.")
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph(("       La operación corresponde a la planilla de {producto} en moneda {moneda} del mes de {mes}.").format(producto=producto, moneda=currency, mes=translate_month_to_spanish(get_current_month())))
    exclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    exclusion_paragraph.runs[0].font.name = 'Arial'

    # Add a table operacion
    operacion_table = doc.add_table(rows=2, cols=8)
    operacion_table.style = 'Table Grid'
    hdr_cells = operacion_table.rows[0].cells
    hdr_cells[0].text = 'Nombre del cliente'
    hdr_cells[1].text = 'Nro. Documento'
    hdr_cells[2].text = 'Fecha Nacimiento'
    hdr_cells[3].text = 'Capital Asegurado'
    hdr_cells[4].text = 'Costo del Seguro'
    hdr_cells[5].text = 'Fecha Vencimiento'
    hdr_cells[6].text = 'Nro. Operación'
    hdr_cells[7].text = 'Motivo'
    for cell in operacion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    row_cells = operacion_table.rows[1].cells
    row_cells[0].text = data_row[0]
    row_cells[1].text = str(data_row[1])
    row_cells[2].text = format_date(data_row[2])
    if currency == "guaraníes":
        row_cells[3].text = str("{:,.0f}".format(data_row[6])).replace(",", ".")
        row_cells[4].text = str("{:,.0f}".format(data_row[8])).replace(",", ".")
    elif currency == "dólares americanos":
        row_cells[3].text = str("{:,.2f}".format(data_row[6])).replace(".", "x").replace(",", ".").replace("x", ",")
        row_cells[4].text = str("{:,.2f}".format(data_row[8])).replace(".", "x").replace(",", ".").replace("x", ",")
    row_cells[5].text = format_date(data_row[10]) 
    row_cells[6].text = str(data_row[4]) 
    row_cells[7].text = " "
    for row in operacion_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'

    # Definir el color gris (RGB: 192, 192, 192)
    gray_color = RGBColor(240, 240, 240)

    # Agregar fondo de color gris a la celda
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}" />')
        tcPr.append(shading_elm)
   
    # Add a section break
    doc.add_paragraph(" ")


def generate_operacion_adelantada_template(doc, data_row, currency, producto):
    from actions.file_actions import translate_month_to_spanish, get_current_month, format_date
    # Add texto exclusion
    exclusion_paragraph = doc.add_paragraph("       Por la presente se informa la exclusión del Prestatario indicado a continuación, de la póliza de Seguro de Vida Colectivo para Cancelación de Deudas, debido a que sus operaciones se han adelantado, es decir, que no corresponden a la planilla de seguros recibida.")
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph(("       La operación corresponde a la planilla de {producto} en moneda {moneda} del mes de {mes}.").format(producto=producto, moneda=currency, mes=translate_month_to_spanish(get_current_month())))
    exclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    exclusion_paragraph.runs[0].font.name = 'Arial'

    # Add a table operacion
    operacion_table = doc.add_table(rows=2, cols=7)
    operacion_table.style = 'Table Grid'
    hdr_cells = operacion_table.rows[0].cells
    hdr_cells[0].text = 'Nombre del cliente'
    hdr_cells[1].text = 'Nro. Documento'
    hdr_cells[2].text = 'Fecha Nacimiento'
    hdr_cells[3].text = 'Capital Asegurado'
    hdr_cells[4].text = 'Costo del Seguro'
    hdr_cells[5].text = 'Fecha de Operación'
    hdr_cells[6].text = 'Nro. Operación'
    for cell in operacion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    row_cells = operacion_table.rows[1].cells
    row_cells[0].text = data_row[0]
    row_cells[1].text = str(data_row[1])
    row_cells[2].text = format_date(data_row[2])
    if currency == "guaraníes":
        row_cells[3].text = str("{:,.0f}".format(data_row[6])).replace(",", ".")
        row_cells[4].text = str("{:,.0f}".format(data_row[8])).replace(",", ".")
    elif currency == "dólares americanos":
        row_cells[3].text = str("{:,.2f}".format(data_row[6])).replace(".", "x").replace(",", ".").replace("x", ",")
        row_cells[4].text = str("{:,.2f}".format(data_row[8])).replace(".", "x").replace(",", ".").replace("x", ",")
    row_cells[5].text = format_date(data_row[9]) 
    row_cells[6].text = str(data_row[4])
    for row in operacion_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'

    # Definir el color gris (RGB: 192, 192, 192)
    gray_color = RGBColor(240, 240, 240)

    # Agregar fondo de color gris a la celda
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}" />')
        tcPr.append(shading_elm)
   
    # Add a section break
    doc.add_paragraph(" ")


def generate_persona_juridica_template(doc, data_row, currency, producto):
    from actions.file_actions import translate_month_to_spanish, get_current_month, format_date
    # Add texto exclusion
    exclusion_paragraph = doc.add_paragraph("       Por la presente se informa la exclusión del Prestatario indicado a continuación, de la póliza de Seguro de Vida Colectivo para Cancelación de Deudas, debido a que es Persona Jurídica.")
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph(("       La operación corresponde a la planilla de {producto} en moneda {moneda} del mes de {mes}.").format(producto=producto, moneda=currency, mes=translate_month_to_spanish(get_current_month())))
    exclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    exclusion_paragraph.runs[0].font.name = 'Arial'

    # Add a table operacion
    operacion_table = doc.add_table(rows=2, cols=8)
    operacion_table.style = 'Table Grid'
    hdr_cells = operacion_table.rows[0].cells
    hdr_cells[0].text = 'Nombre del cliente'
    hdr_cells[1].text = 'Nro. Documento'
    hdr_cells[2].text = 'Fecha Nacimiento'
    hdr_cells[3].text = 'Capital Asegurado'
    hdr_cells[4].text = 'Costo del Seguro'
    hdr_cells[5].text = 'Fecha Inicio'
    hdr_cells[6].text = 'Fecha Vencimiento'
    hdr_cells[7].text = 'Nro. Operación'
    for cell in operacion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    row_cells = operacion_table.rows[1].cells
    row_cells[0].text = data_row[0]
    row_cells[1].text = str(data_row[1])
    row_cells[2].text = format_date(data_row[2])
    if currency == "guaraníes":
        row_cells[3].text = str("{:,.0f}".format(data_row[6])).replace(",", ".")
        row_cells[4].text = str("{:,.0f}".format(data_row[8])).replace(",", ".")
    elif currency == "dólares americanos":
        row_cells[3].text = str("{:,.2f}".format(data_row[6])).replace(".", "x").replace(",", ".").replace("x", ",")
        row_cells[4].text = str("{:,.2f}".format(data_row[8])).replace(".", "x").replace(",", ".").replace("x", ",")
    row_cells[5].text = format_date(data_row[9]) 
    row_cells[6].text = format_date(data_row[10]) 
    row_cells[7].text = str(data_row[4])
    for row in operacion_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'

    # Definir el color gris (RGB: 192, 192, 192)
    gray_color = RGBColor(240, 240, 240)

    # Agregar fondo de color gris a la celda
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}" />')
        tcPr.append(shading_elm)
   
    # Add a section break
    doc.add_paragraph(" ")


def generate_cambio_condiciones_template(doc, data_row, currency, producto):
    from actions.file_actions import translate_month_to_spanish, get_current_month, format_date
    # Add texto exclusion
    exclusion_paragraph = doc.add_paragraph("       Por la presente se informa la exclusión del Prestatario indicado a continuación, de la póliza de Seguro de Vida Colectivo para Cancelación de Deudas, conforme a la Política de Suscripción de la Compañía:")
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph("""
•	La tasa del préstamo asciende a ……………..‰ mensual.
•	Se excluye la cobertura de …………………..
""")
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph(("       La operación corresponde a la planilla de {producto} en moneda {moneda} del mes de {mes}.").format(producto=producto, moneda=currency, mes=translate_month_to_spanish(get_current_month())))
    exclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    exclusion_paragraph.runs[0].font.name = 'Arial'

    # Add a table operacion
    operacion_table = doc.add_table(rows=2, cols=8)
    operacion_table.style = 'Table Grid'
    hdr_cells = operacion_table.rows[0].cells
    hdr_cells[0].text = 'Nombre del cliente'
    hdr_cells[1].text = 'Nro. Documento'
    hdr_cells[2].text = 'Fecha Nacimiento'
    hdr_cells[3].text = 'Capital Asegurado'
    hdr_cells[4].text = 'Costo del Seguro'
    hdr_cells[5].text = 'Fecha Inicio'
    hdr_cells[6].text = 'Fecha Vencimiento'
    hdr_cells[7].text = 'Nro. Operación'
    for cell in operacion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    row_cells = operacion_table.rows[1].cells
    row_cells[0].text = data_row[0]
    row_cells[1].text = str(data_row[1])
    row_cells[2].text = format_date(data_row[2])
    if currency == "guaraníes":
        row_cells[3].text = str("{:,.0f}".format(data_row[6])).replace(",", ".")
        row_cells[4].text = str("{:,.0f}".format(data_row[8])).replace(",", ".")
    elif currency == "dólares americanos":
        row_cells[3].text = str("{:,.2f}".format(data_row[6])).replace(".", "x").replace(",", ".").replace("x", ",")
        row_cells[4].text = str("{:,.2f}".format(data_row[8])).replace(".", "x").replace(",", ".").replace("x", ",")
    row_cells[5].text = format_date(data_row[9]) 
    row_cells[6].text = format_date(data_row[10]) 
    row_cells[7].text = str(data_row[4])
    for row in operacion_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'

    # Definir el color gris (RGB: 192, 192, 192)
    gray_color = RGBColor(240, 240, 240)

    # Agregar fondo de color gris a la celda
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}" />')
        tcPr.append(shading_elm)
   
    # Add a section break
    doc.add_paragraph(" ")


def generate_analisis_template(doc, data_row, currency, producto):
    from actions.file_actions import translate_month_to_spanish, get_current_month, format_date
    # Add texto exclusion
    exclusion_paragraph = doc.add_paragraph("       Por la presente se informa la exclusión del Prestatario indicado a continuación, de la póliza de Seguro de Vida Colectivo para Cancelación de Deudas. Debido a las características del riesgo conforme a la Política de Suscripción de la Compañía, se pospone el análisis del riesgo por 6 meses.")
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph(("       La operación corresponde a la planilla de {producto} en moneda {moneda} del mes de {mes}.").format(producto=producto, moneda=currency, mes=translate_month_to_spanish(get_current_month())))
    exclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    exclusion_paragraph.runs[0].font.name = 'Arial'

    # Add a table operacion
    operacion_table = doc.add_table(rows=2, cols=8)
    operacion_table.style = 'Table Grid'
    hdr_cells = operacion_table.rows[0].cells
    hdr_cells[0].text = 'Nombre del cliente'
    hdr_cells[1].text = 'Nro. Documento'
    hdr_cells[2].text = 'Fecha Nacimiento'
    hdr_cells[3].text = 'Capital Asegurado'
    hdr_cells[4].text = 'Costo del Seguro'
    hdr_cells[5].text = 'Fecha Inicio'
    hdr_cells[6].text = 'Fecha Vencimiento'
    hdr_cells[7].text = 'Nro. Operación'
    for cell in operacion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    row_cells = operacion_table.rows[1].cells
    row_cells[0].text = data_row[0]
    row_cells[1].text = str(data_row[1])
    row_cells[2].text = format_date(data_row[2])
    if currency == "guaraníes":
        row_cells[3].text = str("{:,.0f}".format(data_row[6])).replace(",", ".")
        row_cells[4].text = str("{:,.0f}".format(data_row[8])).replace(",", ".")
    elif currency == "dólares americanos":
        row_cells[3].text = str("{:,.2f}".format(data_row[6])).replace(".", "x").replace(",", ".").replace("x", ",")
        row_cells[4].text = str("{:,.2f}".format(data_row[8])).replace(".", "x").replace(",", ".").replace("x", ",")
    row_cells[5].text = format_date(data_row[9]) 
    row_cells[6].text = format_date(data_row[10]) 
    row_cells[7].text = str(data_row[4])
    for row in operacion_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'

    # Definir el color gris (RGB: 192, 192, 192)
    gray_color = RGBColor(240, 240, 240)

    # Agregar fondo de color gris a la celda
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}" />')
        tcPr.append(shading_elm)
   
    # Add a section break
    doc.add_paragraph(" ")
    despedida_paragraph = doc.add_paragraph("       En tal sentido y hasta la regularización de la Declaración de Salud, el riesgo queda excluido de la póliza. Se adjunta copia del formulario.")
    despedida_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY 
    despedida_paragraph.runs[0].font.name = 'Arial'


def generate_historial_cobertura_template(doc, data_row, currency, producto):
    from actions.file_actions import translate_month_to_spanish, get_current_month, format_date
    # Add texto exclusion
    exclusion_paragraph = doc.add_paragraph("       Por la presente se informa la exclusión del Prestatario indicado a continuación, de la póliza de Seguro de Vida Colectivo para Cancelación de Deudas, debido a que su operación no cuenta con historial de cobertura.")
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph(("       La operación corresponde a la planilla de {producto} en moneda {moneda} del mes de {mes}.").format(producto=producto, moneda=currency, mes=translate_month_to_spanish(get_current_month())))
    exclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    exclusion_paragraph.runs[0].font.name = 'Arial'

    # Add a table operacion
    operacion_table = doc.add_table(rows=2, cols=8)
    operacion_table.style = 'Table Grid'
    hdr_cells = operacion_table.rows[0].cells
    hdr_cells[0].text = 'Nombre del cliente'
    hdr_cells[1].text = 'Nro. Documento'
    hdr_cells[2].text = 'Fecha Nacimiento'
    hdr_cells[3].text = 'Capital Asegurado'
    hdr_cells[4].text = 'Costo del Seguro'
    hdr_cells[5].text = 'Fecha Inicio'
    hdr_cells[6].text = 'Fecha Vencimiento'
    hdr_cells[7].text = 'Nro. Operación'
    for cell in operacion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    row_cells = operacion_table.rows[1].cells
    row_cells[0].text = data_row[0]
    row_cells[1].text = str(data_row[1])
    row_cells[2].text = format_date(data_row[2])
    if currency == "guaraníes":
        row_cells[3].text = str("{:,.0f}".format(data_row[6])).replace(",", ".")
        row_cells[4].text = str("{:,.0f}".format(data_row[8])).replace(",", ".")
    elif currency == "dólares americanos":
        row_cells[3].text = str("{:,.2f}".format(data_row[6])).replace(".", "x").replace(",", ".").replace("x", ",")
        row_cells[4].text = str("{:,.2f}".format(data_row[8])).replace(".", "x").replace(",", ".").replace("x", ",")
    row_cells[5].text = format_date(data_row[9]) 
    row_cells[6].text = format_date(data_row[10]) 
    row_cells[7].text = str(data_row[4])
    for row in operacion_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'

    # Definir el color gris (RGB: 192, 192, 192)
    gray_color = RGBColor(240, 240, 240)

    # Agregar fondo de color gris a la celda
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}" />')
        tcPr.append(shading_elm)
   
    # Add a section break
    doc.add_paragraph(" ")


def generate_no_renovado_template(doc, data_row, currency, producto):
    from actions.file_actions import translate_month_to_spanish, get_current_month, format_date
    # Add texto exclusion
    exclusion_paragraph = doc.add_paragraph("       Por la presente se informa la exclusión del Prestatario indicado a continuación, de la póliza de Seguro de Vida Colectivo para Cancelación de Deudas, debido a que su operación no cuenta con permanencia debido a que no fue remitida para renovar anteriormente.")
    exclusion_paragraph.runs[0].font.name = 'Arial'
    exclusion_paragraph = doc.add_paragraph(("       La operación corresponde a la planilla de {producto} en moneda {moneda} del mes de {mes}.").format(producto=producto, moneda=currency, mes=translate_month_to_spanish(get_current_month())))
    exclusion_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    exclusion_paragraph.runs[0].font.name = 'Arial'

    # Add a table operacion
    operacion_table = doc.add_table(rows=2, cols=8)
    operacion_table.style = 'Table Grid'
    hdr_cells = operacion_table.rows[0].cells
    hdr_cells[0].text = 'Nombre del cliente'
    hdr_cells[1].text = 'Nro. Documento'
    hdr_cells[2].text = 'Fecha Nacimiento'
    hdr_cells[3].text = 'Capital Asegurado'
    hdr_cells[4].text = 'Costo del Seguro'
    hdr_cells[5].text = 'Fecha Inicio'
    hdr_cells[6].text = 'Fecha Vencimiento'
    hdr_cells[7].text = 'Nro. Operación'
    for cell in operacion_table.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
    row_cells = operacion_table.rows[1].cells
    row_cells[0].text = data_row[0]
    row_cells[1].text = str(data_row[1])
    row_cells[2].text = format_date(data_row[2])
    if currency == "guaraníes":
        row_cells[3].text = str("{:,.0f}".format(data_row[6])).replace(",", ".")
        row_cells[4].text = str("{:,.0f}".format(data_row[8])).replace(",", ".")
    elif currency == "dólares americanos":
        row_cells[3].text = str("{:,.2f}".format(data_row[6])).replace(".", "x").replace(",", ".").replace("x", ",")
        row_cells[4].text = str("{:,.2f}".format(data_row[8])).replace(".", "x").replace(",", ".").replace("x", ",")
    row_cells[5].text = format_date(data_row[9]) 
    row_cells[6].text = format_date(data_row[10]) 
    row_cells[7].text = str(data_row[4])
    for row in operacion_table.rows:
        for cell in row.cells:
            # Obtener la fuente actual y establecer el tamaño
            font = cell.paragraphs[0].runs[0].font
            font.size = Pt(9)
            font.name = 'Arial'

    # Definir el color gris (RGB: 192, 192, 192)
    gray_color = RGBColor(240, 240, 240)

    # Agregar fondo de color gris a la celda
    for cell in hdr_cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{gray_color}" />')
        tcPr.append(shading_elm)
   
    # Add a section break
    doc.add_paragraph(" ")


def generate_template_with_content(doc, entity_name, currency, producto, data_row):
    from actions.file_actions import get_formatted_date, get_receptor_segun_entidad
    """Generates a Word file using the appropriate template."""
    # Add date
    date_paragraph = doc.add_paragraph("Encarnación, " + get_formatted_date())
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    date_paragraph.runs[0].font.name = 'Arial'

    # Add entidad
    entidad_paragraph = doc.add_paragraph("""    Señores
    {entity_name}
    Presente""".format(entity_name=entity_name))
    entidad_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    entidad_paragraph.runs[0].font.name = 'Arial'

    # Add receptor
    receptor_paragraph = doc.add_paragraph("""Atn: {receptor}
                        Ref.: Exclusión en Seguro de Vida Cancelación de Deudas 
                        Nota.  N°: /2023""".format(receptor=get_receptor_segun_entidad(entity_name)))
    receptor_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    receptor_paragraph.runs[0].font.name = 'Arial'

    template_name = str(data_row[14])    # Lugar de la lista donde se encuentra el motivo de la exclusion
    if template_name == "CT3":
        generate_cumulo_template(doc, data_row, currency, producto)
    elif template_name == "FF1":
        generate_fallecimiento_template(doc, data_row, currency, producto)
    elif template_name == "DS4":
        generate_falta_ds_template(doc, data_row, currency, producto)
    elif template_name == "DS1":
        generate_ds_incompleta_template(doc, data_row, currency, producto)
    elif template_name == "sin capital":
        generate_sin_capital_template(doc, data_row, currency, producto)
    elif template_name == "MM1":
        generate_mora_template(doc, data_row, currency, producto)
    elif template_name == "ED1":
        generate_edad_template(doc, data_row, currency, producto)
    elif template_name == "OV1":
        generate_operacion_vencida_template(doc, data_row, currency, producto)
    elif template_name == "PS1":
        generate_policita_suscripcion_template(doc, data_row, currency, producto)
    elif template_name == "anulado":
        generate_anulado_template(doc, data_row, currency, producto)
    elif template_name == "DS2":
        generate_diferencia_ds_template(doc, data_row, currency, producto)    
    elif template_name == "DS3":
        generate_informacion_adicional_template(doc, data_row, currency, producto)
    elif template_name == "EP1":
        generate_exclusiones_previas_template(doc, data_row, currency, producto)
    elif template_name == "OA1":
        generate_operacion_adelantada_template(doc, data_row, currency, producto)
    elif template_name == "PJ1":
        generate_persona_juridica_template(doc, data_row, currency, producto)
    elif template_name == "CC1":
        generate_cambio_condiciones_template(doc, data_row, currency, producto)
    elif template_name == "PP1":
        generate_analisis_template(doc, data_row, currency, producto)
    elif template_name == "SH1":
        generate_historial_cobertura_template(doc, data_row, currency, producto)
    elif template_name == "NR1":
        generate_no_renovado_template(doc, data_row, currency, producto)

    # Add despedida
    despedida_paragraph = doc.add_paragraph("       Sin otro particular nos despedimos de usted, atentamente.")
    despedida_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY 
    despedida_paragraph.runs[0].font.name = 'Arial'

    # Save the document
    global file_counter
    template_filename = f"{entity_name}_{template_name}_{currency}_{file_counter}_template.docx"
    file_counter += 1
    doc.save(template_filename)
    return template_filename
