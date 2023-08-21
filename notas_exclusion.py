import openpyxl
from docx import Document

def extract_names_from_excel(excel_file, sheet_name, column_letter=1):
    names = []
    documents = []
    motivos_exclusion = []
    wb = openpyxl.load_workbook(excel_file)
    sheet = wb[sheet_name]
    
    for row in sheet.iter_rows(min_row=2, max_row=sheet.max_row, min_col=column_letter, max_col=3):
        name = row[0].value
        document = row[1].value
        motivo_exclusion = row[2].value
        if name:
            names.append(name)
            documents.append(document)
            motivos_exclusion.append(motivo_exclusion)
    
    return names, documents, motivos_exclusion

def write_names_to_word(exclusion_list, word_file):
    doc = Document()
    doc.add_paragraph("Names extracted from Excel:")
    names, documents, motivos_exclusion = exclusion_list
    
    for name, document, motivo_exclusion in zip(names, documents, motivos_exclusion):
        doc.add_paragraph(note_template(name, document, motivo_exclusion))
    
    doc.save(word_file)

def note_template(name, document, motivo_exclusion):
    return f"The name of the client is {name}, with document {document} and the motive of exclusion is {motivo_exclusion}"

# Define file paths
excel_file_path = 'Test exclusiones.xlsx'
word_file_path = 'Notas exclusiones.docx'
text_template = "The name of the client is [list_element], and the motive of exclusion is"

# Extract names from Excel and write to Word
exclusion_list = extract_names_from_excel(excel_file_path, 'resumen', 1)
write_names_to_word(exclusion_list, word_file = word_file_path)

print("Names extracted from Excel and written to Word.")
